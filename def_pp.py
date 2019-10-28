
import docx


def search_word(file, word, sign):


    file_name = file
    search_word = word
    document_sign = sign

    result_file = open(file_name[:len(file_name)-5]+'-'+search_word+'.txt', 'a')


    doc = docx.Document(file_name)

    print('Jeszcze moment')
    search_index = 0
    for num_line in range(len(doc.paragraphs)):
        if search_word in doc.paragraphs[num_line].text or search_word.capitalize() in doc.paragraphs[num_line].text:
            step = 0
            while(True):
                if document_sign in doc.paragraphs[num_line-step].text:
                    if step != 0:
                        result_file.write(doc.paragraphs[num_line-step].text[:30]+'...')
                        result_file.write('\n')
                    result_file.write(doc.paragraphs[num_line].text)
                    result_file.write('\n')
                    result_file.write('\n')
                    search_index += 1
                    print(f'Znaleziono {search_index} ')
                    break
                else:
                    step += 1

    result_file.close()