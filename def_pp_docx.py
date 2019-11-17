
import docx
#zapisuje do pliku docx

def search_word(file, word, sign):

    def delete_paragraph(paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None

    file_name = file
    search_word = word
    document_sign = sign

    result_file = docx.Document()

    doc = docx.Document(file_name)
    print('Jeszcze moment')
    kropka = 0
    kasuj = 'MontyPython:)'


    for num_line in range(len(doc.paragraphs)):#dla każdego paragrafu po kolei
        find = False
        search_index = 0
        if len(doc.paragraphs[num_line].text) > 5:
            if document_sign in doc.paragraphs[num_line].text[:5]:#jeżeli w paragrafie jest 'Art.' lub '§'
                search_index = num_line + 1
                if len(doc.paragraphs[search_index].text)>5:
                    while(document_sign not in doc.paragraphs[search_index].text[:5]): # to w nastepnych paragrafach szukaj słowa
                        warunek1 = search_word not in doc.paragraphs[search_index].text
                        warunek2 = search_word.capitalize() not in doc.paragraphs[search_index].text
                        if (warunek1 and warunek2):
                            doc.paragraphs[search_index].text = '' # jezeli nie znajdziesz, wykasuj paragraf
                        else:
                            find = True #znacznik czy znaleziono szukane słowo
                            print('.', end='')
                            kropka += 1
                            if kropka % 100 == 0:
                                print('\n')
                        if search_index + 1 < len(doc.paragraphs):
                            search_index += 1
                        else:
                            break
                    warunek1 = search_word in doc.paragraphs[num_line].text
                    warunek2 = search_word.capitalize() in doc.paragraphs[num_line].text
                    if (not find and not warunek1 and not warunek2):
                        doc.paragraphs[num_line].text = ''

    num_line = 0
    kropka = 0
    print('kasowanie..')
    while(num_line < len(doc.paragraphs)):  # dla każdego paragrafu po kolei
        if (doc.paragraphs[num_line].text == ''):
            delete_paragraph(doc.paragraphs[num_line])
            print('.', end='')
            kropka += 1
            if kropka % 100 == 0:
                print('\n')
        else:
            num_line += 1

    doc.save(file_name[:len(file_name)-5]+'-'+search_word+'.docx')

